from django.shortcuts import render, get_object_or_404, redirect
from django.core.mail import send_mail
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.contrib.auth.decorators import login_required
from django.contrib.auth import authenticate, login, logout
from django.conf import settings
from .models import Offerte 
import json
from django.views.decorators.csrf import ensure_csrf_cookie
import requests
import base64
from docx import Document
from docx.shared import Pt
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from django.http import HttpResponse
from .forms import ContactForm 
from .forms import OfferteForm
from django.contrib import messages
from .forms import CateringRequestForm
from .models import Note, Klant

from django.views.decorators.http import require_POST
from decimal import Decimal
from datetime import date, timedelta
from django.views.decorators.csrf import csrf_protect
from django.views.generic import ListView
from .models import ActieVanHetMoment
from django.utils.timezone import now  # Voor de today datum
from django.utils.html import strip_tags

import re
from django.http import JsonResponse
from django.views.decorators.csrf import csrf_exempt
from django.core.mail import send_mail
from django.http import HttpResponse
import xlsxwriter
import io



EMAIL_HOST_USER = settings.EMAIL_HOST_USER





def index(request):
    acties = ActieVanHetMoment.objects.filter(
        actief_tot__gte=date.today()
    ).order_by('-actief_tot')[:5]  # Toon alleen actieve acties, max 3
    
    context = {
        'acties': acties,
        # ... je bestaande context
    }
    return render(request, 'index.html', context)

def contact(request):
    acties = ActieVanHetMoment.objects.filter(
        actief_tot__gte=date.today()
    ).order_by('-actief_tot')[:5]  # Toon alleen actieve acties, max 3
    
    context = {
        'acties': acties,
        # ... je bestaande context
    }
    return render(request, 'contact.html', context)

def formules(request):
    acties = ActieVanHetMoment.objects.filter(
        actief_tot__gte=date.today()
    ).order_by('-actief_tot')[:5]  # Toon alleen actieve acties, max 3
    
    context = {
        'acties': acties,
        # ... je bestaande context
    }
    return render(request, 'formules.html', context)

def offerte(request):
    acties = ActieVanHetMoment.objects.filter(
        actief_tot__gte=date.today()
    ).order_by('-actief_tot')[:5]  # Toon alleen actieve acties, max 3
    
    context = {
        'acties': acties,
        'openroute_api_key': settings.OPENROUTE_API_KEY,  # Add this line
        # ... je bestaande context
    }
    return render(request, 'offerte.html', context)

def snackssauzen(request): 
    return render(request, 'snackssauzen.html')

def faq(request):  
    return render(request, 'faq.html')

def voorwaarden(request):
    return render(request, 'voorwaarden.html')

    
def privacy(request):
    return render(request, 'Privacybeleid.html')

def event(request):
    acties = ActieVanHetMoment.objects.filter(
        actief_tot__gte=date.today()
    ).order_by('-actief_tot')[:5]  # Toon alleen actieve acties, max 3
    
    context = {
        'acties': acties,
        # ... je bestaande context
    }
    return render(request, 'event.html', context)

def klant_aanpassen(request, klant_id):

    # Haal de klant op basis van het ID op uit de URL parameter
    klant = get_object_or_404(Klant, id=klant_id)
    
    if request.method == 'POST':
        # Verwerk het formulier
        klant.naam_bedrijf = request.POST.get('naam_bedrijf')
        klant.naam_contactpersoon = request.POST.get('naam_contactpersoon')
        klant.email = request.POST.get('email')
        klant.straat = request.POST.get('straat')
        klant.nummer = request.POST.get('nummer')
        klant.gemeente = request.POST.get('gemeente')
        klant.telefoon = request.POST.get('telefoon')
        klant.btw_nummer = request.POST.get('btw_nummer')
        
        # Sla de wijzigingen op
        klant.save()
        
        # Voeg een succesmelding toe
        messages.success(request, f"Klant '{klant.naam_bedrijf}' is succesvol bijgewerkt.")
        
        # Redirect naar de klantenpagina
        return redirect('klanten')
    
    # Geef het formulier weer met de huidige klantgegevens
    return render(request, 'klant_aanpassen.html', {'klant': klant})


# alle events die ingeplandt zijn op de page geplande_feesten.html
@login_required
def geplande_feesten(request):
    offertes = Offerte.objects.filter(is_bevestigd="Feest ingepland").order_by('gewenste_datum', 'gewenste_tijd')
    return render(request, 'geplande_feesten.html', {'offertes': offertes})

# offertes maken op het dashboard -------------------------------------------------------------
@login_required
def zelf_offerte(request):
    if request.method == 'POST':
        try:
            # Debug: Print all POST data
            print("POST data:", request.POST)
            
            # Get form data with validation
            naam_contactpersoon = request.POST.get('naam_contactpersoon')
            email = request.POST.get('email')
            aantal_personen = int(request.POST.get('aantal_personen', 0))
            straat = request.POST.get('straat')
            nummer = request.POST.get('nummer')
            gemeente = request.POST.get('gemeente')
            event_adres = request.POST.get('event_adres')
            telefoon = request.POST.get('telefoon')
            naam_bedrijf = request.POST.get('naam_bedrijf', '')
            btw_nummer = request.POST.get('btw_nummer', '')
            
            # Date and time handling
            gewenste_datum = request.POST.get('gewenste_datum')
            gewenste_tijd = request.POST.get('gewenste_tijd')
            
            # Haal de formule naam op uit het nieuwe verborgen veld
            formule_naam = request.POST.get('formule_naam', '')
            
            # Als er geen formule_naam is, gebruik dan de waarde om de juiste tekst te bepalen
            if not formule_naam:
                formule_waarde = request.POST.get('formule')
                if formule_waarde == '8':
                    formule_naam = '⭐ Formule 1'
                elif formule_waarde == '9':
                    formule_naam = '⭐⭐ Formule 2'
                elif formule_waarde == '10.5':
                    formule_naam = '⭐⭐⭐ Luxe Formule'
                else:
                    formule_naam = f'Formule (€{formule_waarde})'
            
            extra_info = request.POST.get('extra_info', '')
            
            # Numeric fields with error checking
            try:
                afstand = float(request.POST.get('afstand', 0))
            except ValueError:
                afstand = 0
                
            try:
                basis_prijs = Decimal(request.POST.get('basis_prijs', 0))
            except (ValueError, TypeError):
                basis_prijs = Decimal('0')
                
            try:
                extra_kosten = Decimal(request.POST.get('extra_kosten', 0) or 0)
            except (ValueError, TypeError):
                extra_kosten = Decimal('0')
                
            try:
                korting = Decimal(request.POST.get('korting', 0) or 0)
            except (ValueError, TypeError):
                korting = Decimal('0')
                
            try:
                totaal_prijs = Decimal(request.POST.get('totaal_prijs', 0))
            except (ValueError, TypeError):
                # Calculate total price if not provided correctly
                totaal_prijs = basis_prijs + extra_kosten - korting
                
            is_bevestigd = request.POST.get('is_bevestigd', 'Nieuwe offerte')
            
            # Debug: Print processed values
            print(f"Processed data: naam={naam_contactpersoon}, formule={formule_naam}, prijs={totaal_prijs}")
            
            # Create and save new offerte
            offerte = Offerte(
                naam_contactpersoon=naam_contactpersoon,
                email=email,
                aantal_personen=aantal_personen,
                straat=straat,
                nummer=nummer,
                gemeente=gemeente,
                event_adres=event_adres,
                telefoon=telefoon,
                naam_bedrijf=naam_bedrijf,
                btw_nummer=btw_nummer,
                gewenste_datum=gewenste_datum,
                gewenste_tijd=gewenste_tijd,
                formule=formule_naam,  # Gebruik hier de formule naam in plaats van de waarde
                extra_info=extra_info,
                afstand=afstand,
                extra_kosten=extra_kosten,
                basis_prijs=basis_prijs,
                totaal_prijs=totaal_prijs,
                korting=korting,
                is_bevestigd=is_bevestigd
            )
            
            # Save the offerte
            offerte.save()
            print(f"Offerte saved with ID: {offerte.id}")
            
            # Show success message with specific tag for admin actions
            messages.success(request, f"Nieuwe offerte aangemaakt voor: {naam_bedrijf or naam_contactpersoon}", extra_tags='admin_actie')
            
            # Redirect to quotes overview
            return redirect('offertes')
            
        except Exception as e:
            # Show detailed error
            import traceback
            print(f"Error: {str(e)}")
            print(traceback.format_exc())
            messages.error(request, f"Er is een fout opgetreden: {str(e)}", extra_tags='admin_actie')
    
    # First time loading the page (GET request)
    return render(request, 'zelf_offerte.html')


# offerte aanpassen in de overzicht offertes
@login_required
def offerte_aanpassen(request, offerte_id):
    offerte = get_object_or_404(Offerte, id=offerte_id)
    
    if request.method == "POST":
        # Update alleen velden als er een nieuwe waarde is opgegeven
        if request.POST.get("naam_contactpersoon"):
            offerte.naam_contactpersoon = request.POST.get("naam_contactpersoon")
        if request.POST.get("email"):
            offerte.email = request.POST.get("email")
        if request.POST.get("telefoon"):
            offerte.telefoon = request.POST.get("telefoon")
        if request.POST.get("aantal_personen"):
            offerte.aantal_personen = request.POST.get("aantal_personen")
        if request.POST.get("straat"):
            offerte.straat = request.POST.get("straat")
        if request.POST.get("nummer"):
            offerte.nummer = request.POST.get("nummer")
        if request.POST.get("gemeente"):
            offerte.gemeente = request.POST.get("gemeente")
        if request.POST.get("event_adres"):
            offerte.event_adres = request.POST.get("event_adres")
        if request.POST.get("btw_nummer"):
            offerte.btw_nummer = request.POST.get("btw_nummer")
        if request.POST.get("gewenste_datum"):
            offerte.gewenste_datum = request.POST.get("gewenste_datum")
        # Alleen gewenste_tijd bijwerken als er een niet-lege waarde is
        gewenste_tijd = request.POST.get("gewenste_tijd")
        if gewenste_tijd:  # Alleen updaten als er een waarde is opgegeven
            offerte.gewenste_tijd = gewenste_tijd
        # Als gewenste_tijd leeg is, blijft de bestaande waarde behouden (geen actie nodig)
        if request.POST.get("formule"):
            offerte.formule = request.POST.get("formule")
        if request.POST.get("extra_info"):
            offerte.extra_info = request.POST.get("extra_info")
        if request.POST.get("afstand"):
            offerte.afstand = request.POST.get("afstand")
        if request.POST.get("extra_kosten"):
            offerte.extra_kosten = request.POST.get("extra_kosten")
        if request.POST.get("basis_prijs"):
            offerte.basis_prijs = request.POST.get("basis_prijs")
        if request.POST.get("totaal_prijs"):
            offerte.totaal_prijs = request.POST.get("totaal_prijs")
        if request.POST.get("korting"):
            offerte.korting = request.POST.get("korting")
        offerte.save()
        return redirect("offertes_overzicht")
    
    return render(request, "offerte_aanpassen.html", {"offerte": offerte})

@login_required
def offertes_overzicht(request):
    offertes = Offerte.objects.all()
    return render(request, "offertes_overzicht.html", {"offertes": offertes})

@login_required
def klanten_view(request):
    klanten = Klant.objects.all()
    return render(request, "klanten.html", {"klanten": klanten})



# gewoon contactform contactform
def contact(request):
    if request.method == 'POST':
        form = ContactForm(request.POST)
        if form.is_valid():
            name = form.cleaned_data['name']
            email = form.cleaned_data['email']
            subject = form.cleaned_data['subject']
            message = form.cleaned_data['message']

            # --- Maak de HTML e-mail in de gewenste stijl ---
            html_message = f"""
            <html>
            <head>
                <style>
                    body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                    .container {{ width: 80%; margin: 20px auto; border: 1px solid #ddd; padding: 20px; border-radius: 8px; background-color: #f9f9f9; }}
                    h2 {{ color: #0056b3; border-bottom: 2px solid #0056b3; padding-bottom: 10px; margin-bottom: 20px; }}
                    h3 {{ color: #0056b3; margin-bottom: 15px; }}
                    p {{ margin-bottom: 10px; }}
                    strong {{ color: #0056b3; }}
                    ul {{ list-style: none; padding: 0; }}
                    ul li {{ margin-bottom: 5px; }}
                    a {{ color: #007bff; text-decoration: none; }}
                    a:hover {{ text-decoration: underline; }}
                    .info-section {{ background-color: #e6f2ff; padding: 15px; border-radius: 5px; margin-bottom: 15px; }}
                    .message-section {{ background-color: #fff; border: 1px solid #ccc; padding: 15px; border-radius: 5px; }}
                </style>
            </head>
            <body>
                <div class="container">
                    <h2>Contactformulier aanvraag</h2>
                    <div class="info-section">
                        <p><strong>Onderwerp:</strong> {subject}</p>
                        <p><strong>Naam:</strong> {name}</p>
                        <p><strong>E-mail:</strong> <a href="mailto:{email}">{email}</a></p>
                    </div>
                    <div class="message-section">
                        <p><strong>Bericht:</strong></p>
                        <p>{message}</p>
                    </div>
                </div>
            </body>
            </html>
            """

            # Maak ook een plain text versie vanuit de HTML
            plain_message = strip_tags(html_message)

            try:
                send_mail(
                    subject=f"Contact Formulier: {subject}",
                    message=plain_message, # Gebruik de plain_message die we uit de HTML hebben gehaald
                    from_email=settings.EMAIL_HOST_USER,
                    recipient_list=['Info@retrosnacks.be'],
                    html_message=html_message, # HTML-versie
                    fail_silently=False,
                )
                messages.success(request, 'Uw bericht is verzonden!')
                return redirect('contact')
            except Exception as e:
                messages.error(request, f'Er is een fout opgetreden: {str(e)}')
        else:
            # Als het formulier niet geldig is, render dan de pagina opnieuw met de fouten
            messages.error(request, 'Controleer de ingevoerde gegevens. Er zijn fouten in het formulier.')
    else:
        form = ContactForm()

    return render(request, 'contact.html', {'form': form})



# Groot event contactform
def event(request):
    if request.method == 'POST':
        # Haal alle formuliergegevens op
        name = request.POST.get('name')
        company = request.POST.get('company')
        email = request.POST.get('email')
        phone = request.POST.get('phone')
        event_name = request.POST.get('event_name')
        event_locatie = request.POST.get('event_locatie')
        event_date = request.POST.get('event_date')
        visitors = request.POST.get('visitors')
        foodtrucks = request.POST.getlist('foodtrucks')  # Krijg alle geselecteerde foodtrucks
        extra_info = request.POST.get('extra_info')

        # --- Maak de HTML e-mail ---
        html_message = f"""
        <html>
        <head>
            <style>
                body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                .container {{ width: 80%; margin: 20px auto; border: 1px solid #ddd; padding: 20px; border-radius: 8px; background-color: #f9f9f9; }}
                h2 {{ color: #0056b3; border-bottom: 2px solid #0056b3; padding-bottom: 10px; margin-bottom: 20px; }}
                h3 {{ color: #0056b3; margin-bottom: 15px; }}
                p {{ margin-bottom: 10px; }}
                strong {{ color: #0056b3; }}
                ul {{ list-style: none; padding: 0; }}
                ul li {{ margin-bottom: 5px; }}
                a {{ color: #007bff; text-decoration: none; }}
                a:hover {{ text-decoration: underline; }}
                .info-section {{ background-color: #e6f2ff; padding: 15px; border-radius: 5px; margin-bottom: 15px; }}
                .extra-info {{ background-color: #fff; border: 1px solid #ccc; padding: 15px; border-radius: 5px; }}
            </style>
        </head>
        <body>
            <div class="container">
                <h2>Groot Evenementen Aanvraag</h2>
                <div class="info-section">
                    <p><strong>Contactgegevens:</strong></p>
                    <ul>
                        <li><strong>Naam:</strong> {name}</li>
                        <li><strong>Bedrijf/Organisatie:</strong> {company or 'Niet opgegeven'}</li>
                        <li><strong>E-mail:</strong> <a href="mailto:{email}">{email}</a></li>
                        <li><strong>Telefoonnummer:</strong> {phone}</li>
                    </ul>
                </div>
                <div class="info-section">
                    <p><strong>Evenement Details:</strong></p>
                    <ul>
                        <li><strong>Naam evenement:</strong> {event_name or 'Niet opgegeven'}</li>
                        <li><strong>Datum evenement:</strong> {event_date}</li>
                        <li><strong>Adres evenement:</strong> {event_locatie}</li>
                        <li><strong>Verwacht aantal bezoekers:</strong> {visitors}</li>
                        <li><strong>Gewenste foodtruck(s):</strong> {', '.join(foodtrucks) if foodtrucks else 'Geen geselecteerd'}</li>
                    </ul>
                </div>
                <div class="extra-info">
                    <p><strong>Extra informatie:</strong></p>
                    <p>{extra_info or 'Geen extra info'}</p>
                </div>
            </div>
        </body>
        </html>
        """

        # Maak ook een plain text versie (altijd goed voor als HTML niet werkt)
        plain_message = strip_tags(html_message) # Genereer plain text vanuit HTML

        try:
            # Verstuur de e-mail
            send_mail(
                subject=f"Catering Aanvraag: {event_date} / {visitors} personen",
                message=plain_message, # Platte tekstversie
                from_email=settings.EMAIL_HOST_USER,
                recipient_list=['Info@retrosnacks.be'],
                html_message=html_message, # HTML-versie
                fail_silently=False,
            )
            messages.success(request, 'Uw aanvraag is succesvol verzonden!')
            return redirect('event')  # Herlaad de pagina na succes
        except Exception as e:
            messages.error(request, f'Er is een fout opgetreden: {str(e)}')

    # Render hetzelfde template of een bedankpagina
    return render(request, 'event.html')







# Facturen bekijk -----------------------------------------------------------------------------------------
@login_required
def invoice_list(request):
    url = "https://eenvoudigfactureren.be/api/v1/invoices"
    headers = {
        "Accept": "application/json"
    }
    auth = (settings.EENVOUDIGFACTUREREN_EMAIL, settings.EENVOUDIGFACTUREREN_PASSWORD)

    try:
        response = requests.get(url, headers=headers, auth=auth)
        response.raise_for_status()  # Gooit een uitzondering bij HTTP-fouten
        invoices = response.json()
    except requests.exceptions.RequestException as e:
        print(f"Fout bij het ophalen van facturen: {e}")
        invoices = []

    return render(request, 'facturen.html', {'invoices': invoices})








def login_page(request):  # Toont alleen de login-pagina--------------------------------------------------------
    return render(request, 'login.html')


def login_view(request):  # Verwerkt de inloglogica
    if request.method == 'POST':
        username = request.POST.get('username')
        password = request.POST.get('password')
        user = authenticate(request, username=username, password=password)
        if user is not None:
            login(request, user)
            return redirect('dashboard')
        else:
            return render(request, 'login.html', {'error': 'Invalid credentials'})
    return render(request, 'login.html')

from django.contrib.auth.decorators import login_required
from django.shortcuts import render, redirect
from datetime import date, timedelta
from .models import Note, Offerte, ActieVanHetMoment
from .forms import ActieForm

from django.shortcuts import render, redirect, get_object_or_404
from django.contrib.auth.decorators import login_required
from datetime import date, timedelta

from .models import ActieVanHetMoment, Note, Offerte # Zorg dat al je modellen hier zijn geïmporteerd
from .forms import ActieForm # Jouw formulier voor ActieVanHetMoment

@login_required
def dashboard_view(request):
    if request.method == 'POST':
        # Verwerk formulier voor 'Actie van het Moment'
        if 'actie_submit' in request.POST:
            actie_form = ActieForm(request.POST)
            if actie_form.is_valid():
                actie_form.save()
                return redirect('dashboard')
        
        # Verwerk verwijderen van 'Actie van het Moment'
        if 'delete_actie_submit' in request.POST:
            actie_id = request.POST.get('actie_id')
            if actie_id:
                # Gebruik get_object_or_404 voor robuustheid en veiligheid
                # Zorg dat alleen beheerders/juiste gebruikers dit kunnen, indien nodig.
                # Voor nu ga ik ervan uit dat login_required volstaat.
                actie_te_verwijderen = get_object_or_404(ActieVanHetMoment, pk=actie_id)
                actie_te_verwijderen.delete()
                return redirect('dashboard')

        # Andere acties (notes)
        action = request.POST.get('action', '')
        if action == 'add_note':
            title = request.POST.get('title')
            content = request.POST.get('content')
            if title and content:
                Note.objects.create(user=request.user, title=title, content=content)
        elif action == 'delete_note':
            note_id = request.POST.get('note_id')
            if note_id:
                Note.objects.filter(id=note_id, user=request.user).delete()
        elif action == 'edit_note':
            note_id = request.POST.get('note_id')
            title = request.POST.get('title')
            content = request.POST.get('content')
            if note_id and title and content:
                try:
                    note = Note.objects.get(id=note_id, user=request.user)
                    note.title = title
                    note.content = content
                    note.save()
                except Note.DoesNotExist:
                    pass
        
        # Na het verwerken van elke POST actie, redirect om dubbele submits te voorkomen
        return redirect('dashboard')

    # Voor GET-verzoeken of na een redirect
    actie_form = ActieForm() # Initialiseer het formulier voor GET-verzoeken

    # Data ophalen
    notes = Note.objects.filter(user=request.user).order_by('-created_at')
    acties = ActieVanHetMoment.objects.all().order_by('-actief_tot') # Let op: ik heb de modelnaam aangepast naar ActieVanHetMoment

    today = date.today()
    start_week = today - timedelta(days=today.weekday())
    end_week = start_week + timedelta(days=6)
    week_events = Offerte.objects.filter(
        gewenste_datum__range=(start_week, end_week),
        is_bevestigd="Feest ingepland"
    ).order_by('gewenste_datum', 'gewenste_tijd')

    context = {
        'notes': notes,
        'week_events': week_events,
        'acties': acties,
        'actie_form': actie_form,
    }

    return render(request, 'dashboard.html', context)


def logout_view(request):  # Logout functionaliteit
    logout(request)
    return redirect('login')

@login_required
def offertes_view(request):
    offertes = Offerte.objects.all().order_by('-created_at')  # Nieuwste offertes eerst
    return render(request, 'offertes.html', {'offertes': offertes})

@login_required
def facturen_view(request):
    return render(request, 'facturen.html')



# Klant ver<ijderen in klanten.html ------------------------------------------------------------
@login_required
def klant_verwijderen(request, klant_id):
    if request.method == 'POST':
        klant = get_object_or_404(Klant, id=klant_id)
        klant.delete()
        messages.success(request, 'Klant is succesvol verwijderd.', extra_tags='admin_actie')
    return redirect('klanten')  # Pas aan naar jouw viewnaam voor de klantenlijst

# alles voor de offerte maak page ------------------------------------------------------------
@require_POST
@csrf_protect
def save_offerte(request):
    if request.method == "POST":
        try:
            # Probeer JSON data te laden als de Content-Type application/json is
            if 'application/json' in request.content_type:
                data = json.loads(request.body)
            else:
                # Als het form-data is (bijvoorbeeld van een standaard HTML form post),
                # dan is het request.POST al een dictionary-achtig object.
                # Voor consistentie maken we hier een kopie.
                data = request.POST.copy()

            # Normaliseer gewenste_tijd voordat het aan het formulier wordt doorgegeven
            gewenste_tijd_raw = data.get("gewenste_tijd", "").strip()
            print(f"Ontvangen tijd (raw): '{gewenste_tijd_raw}'")

            try:
                gewenste_tijd_raw = gewenste_tijd_raw.replace('.', ':')
                if re.match(r'^\d{1,2}:\d{1,2}(:\d{1,2})?$', gewenste_tijd_raw):
                    parts = gewenste_tijd_raw.split(':')
                    hours = int(parts[0])
                    minutes = int(parts[1])
                    gewenste_tijd_norm = f"{hours:02d}:{minutes:02d}"
                else:
                    try:
                        tijd_float = float(gewenste_tijd_raw)
                        hours = int(tijd_float)
                        minutes = int((tijd_float - hours) * 60)
                        gewenste_tijd_norm = f"{hours:02d}:{minutes:02d}"
                    except ValueError:
                        print(f"Kon tijd niet parsen: '{gewenste_tijd_raw}', standaard op 00:00")
                        gewenste_tijd_norm = "00:00" # Of laat validatie dit afhandelen
            except Exception as e:
                print(f"Fout bij verwerken tijd: {str(e)}")
                gewenste_tijd_norm = "00:00" # Of laat validatie dit afhandelen

            data['gewenste_tijd'] = gewenste_tijd_norm # Update de data dictionary

            # Maak een instantie van het formulier met de ontvangen data
            form = OfferteForm(data)

            if form.is_valid():
                cleaned_data = form.cleaned_data
                print(f"Formulier data is valide: {cleaned_data}")

                # Maak en sla de Offerte op met cleaned_data
                offerte = Offerte(
                    naam_contactpersoon=cleaned_data["naam_contactpersoon"],
                    email=cleaned_data["email"],
                    aantal_personen=cleaned_data["aantal_personen"],
                    straat=cleaned_data["straat"],
                    nummer=cleaned_data["nummer"],
                    gemeente=cleaned_data["gemeente"],
                    event_adres=data.get("event_adres", f"{cleaned_data['straat']} {cleaned_data['nummer']}, {cleaned_data['gemeente']}"), # Fallback
                    telefoon=cleaned_data["telefoon"],
                    naam_bedrijf=cleaned_data.get("naam_bedrijf", ""),
                    btw_nummer=cleaned_data.get("btw", ""), # Let op: in form is het 'btw'
                    gewenste_datum=cleaned_data["gewenste_datum"],
                    gewenste_tijd=cleaned_data["gewenste_tijd"],
                    formule=cleaned_data["formule"],
                    extra_info=cleaned_data.get("extra_info", ""),
                    afstand=cleaned_data.get("afstand"), # Komen uit JS, niet direct uit formulier invulvelden
                    extra_kosten=cleaned_data.get("extra_kosten"),
                    basis_prijs=cleaned_data.get("basis_prijs"),
                    totaal_prijs=cleaned_data.get("totaal_prijs"),
                    korting=cleaned_data.get("korting"),
                    is_bevestigd="Nieuwe offerte"
                )
                offerte.save()

                # E-mail logica (nu met HTML-opmaak)
                from_email = settings.EMAIL_HOST_USER
                owner_email = "Info@retrosnacks.be" # Verander indien nodig

                # --- E-mail voor de eigenaar (met opmaak) ---
                owner_subject = f"Nieuwe offerte aanvraag van {cleaned_data['naam_contactpersoon']}"
                owner_html_message = f"""
                <html>
                <head>
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                        .container {{ width: 80%; margin: 20px auto; border: 1px solid #ddd; padding: 20px; border-radius: 8px; background-color: #f9f9f9; }}
                        h2 {{ color: #0056b3; border-bottom: 2px solid #0056b3; padding-bottom: 10px; margin-bottom: 20px; }}
                        p {{ margin-bottom: 10px; }}
                        strong {{ color: #0056b3; }}
                        .section {{ background-color: #eaedf0; padding: 15px; border-radius: 5px; margin-bottom: 15px; }}
                        ul {{ list-style: none; padding: 0; }}
                        ul li {{ margin-bottom: 5px; }}
                        .price-details {{ background-color: #fff; border: 1px solid #ccc; padding: 15px; border-radius: 5px; }}
                        .total-price {{ font-size: 1.2em; font-weight: bold; color: #d9534f; }}
                    </style>
                </head>
                <body>
                    <div class="container">
                        <h2>Nieuwe offerte aanvraag</h2>
                        <div class="section">
                            <p><strong><u>Contactgegevens:</u></strong></p>
                            <ul>
                                <li><strong>Naam contactpersoon:</strong> {cleaned_data['naam_contactpersoon']}</li>
                                <li><strong>Bedrijf:</strong> {cleaned_data.get('naam_bedrijf', '-')}</li>
                                <li><strong>Email:</strong> {cleaned_data['email']}</li>
                                <li><strong>Telefoon:</strong> {cleaned_data['telefoon']}</li>
                                <li><strong>BTW-nummer:</strong> {cleaned_data.get('btw', '-')}</li>
                            </ul>
                        </div>
                        <div class="section">
                            <p><strong><u>Event details:</u></strong></p>
                            <ul>
                                <li><strong>Datum:</strong> {cleaned_data['gewenste_datum']}</li>
                                <li><strong>Begin tijd bakken:</strong> {cleaned_data['gewenste_tijd']}</li>
                                <li><strong>Locatie:</strong> {offerte.event_adres}</li>
                                <li><strong>Aantal personen:</strong> {cleaned_data['aantal_personen']}</li>
                                <li><strong>Formule:</strong> {cleaned_data['formule']}</li>
                            </ul>
                        </div>
                        <div class="price-details">
                            <p><strong><u>Prijs berekening:</u></strong></p>
                            <ul>
                                <li><strong>Basis prijs:</strong> &euro;{float(cleaned_data.get('basis_prijs', 0)):.2f}</li>
                                <li><strong>Extra kosten (afstand {float(cleaned_data.get('afstand', 0)):.1f} km):</strong> &euro;{float(cleaned_data.get('extra_kosten', 0)):.2f}</li>
                                <li><strong>Korting:</strong> &euro;{float(cleaned_data.get('korting', 0)):.2f}</li>
                                <li class="total-price"><strong>Totaal:</strong> &euro;{float(cleaned_data.get('totaal_prijs', 0)):.2f}</li>
                            </ul>
                        </div>
                        <p><strong>Extra info:</strong> {cleaned_data.get('extra_info', '-')}</p>
                    </div>
                </body>
                </html>
                """
                # Voor de platte tekst versie, strip de HTML tags (optioneel, maar goed voor e-mail clients die geen HTML ondersteunen)
                from django.utils.html import strip_tags
                owner_plain_message = strip_tags(owner_html_message)


                # --- E-mail voor de klant (met opmaak) ---
                klant_subject = "Offerte Retrosnacks / Aanvraag beschikbaarheid datum"
                klant_html_message = f"""
                <html>
                <head>
                    <style>
                        body {{ font-family: Arial, sans-serif; line-height: 1.6; color: #333; }}
                        .container {{ width: 80%; margin: 20px auto; border: 1px solid #ddd; padding: 20px; border-radius: 8px; background-color: #f9f9f9; }}
                        /* Groen vervangen door blauw */
                        h2 {{ color: #007bff; border-bottom: 2px solid #007bff; padding-bottom: 10px; margin-bottom: 20px; }}
                        p {{ margin-bottom: 10px; }}
                        /* Groen vervangen door blauw */
                        strong {{ color: #007bff; }}
                        ul {{ list-style: none; padding: 0; }}
                        ul li {{ margin-bottom: 5px; }}
                        .note {{ background-color: #fff3cd; color: #856404; padding: 10px; border: 1px solid #ffeeba; border-radius: 5px; margin-top: 20px; }}
                        .total-price {{ font-size: 1.2em; font-weight: bold; color: #d9534f; }}
                    </style>
                </head>
                <body>
                    <div class="container">
                        <h2>Bedankt voor uw aanvraag bij Retrosnacks!</h2>
                        <p>Beste <strong>{cleaned_data['naam_contactpersoon']}</strong>,</p>
                        <p class="note">Gelieve niet op deze mail te antwoorden, dit is een automatische bevestiging.</p>
                        <p>We hebben uw aanvraag ontvangen en zullen <strong>binnen 24 uur</strong> contact met u opnemen over de beschikbaarheid en eventuele verdere details.</p>

                        <h3>Details van uw aanvraag:</h3>
                        <ul>
                            <li><strong>Datum:</strong> {cleaned_data['gewenste_datum']}</li>
                            <li><strong>Begin tijd bakken:</strong> {cleaned_data['gewenste_tijd']}</li>
                            <li><strong>Locatie:</strong> {offerte.event_adres}</li>
                            <li><strong>Aantal personen:</strong> {cleaned_data['aantal_personen']}</li>
                            <li><strong>Formule:</strong> {cleaned_data['formule']}</li>
                            <li><strong>Extra informatie:</strong> {cleaned_data.get('extra_info', '-')}</li>
                            <li class="total-price"><strong>Totaal bedrag:</strong> &euro;{float(cleaned_data.get('totaal_prijs', 0)):.2f} (Prijs is inclusief 12% BTW en transportkosten)</li>
                        </ul>

                        <p>Met vriendelijke groeten,</p>
                        <p><strong>Het team van Retrosnacks</strong></p>
                    </div>
                </body>
                </html>
                """
                klant_plain_message = strip_tags(klant_html_message)

                try:
                    # Verstuur e-mail naar de eigenaar
                    send_mail(
                        owner_subject,
                        owner_plain_message, # Plain text fallback
                        from_email,
                        [owner_email],
                        html_message=owner_html_message, # HTML-versie
                        fail_silently=False
                    )

                    # Verstuur e-mail naar de klant
                    send_mail(
                        klant_subject,
                        klant_plain_message, # Plain text fallback
                        from_email,
                        [cleaned_data['email']],
                        html_message=klant_html_message, # HTML-versie
                        fail_silently=False
                    )

                except Exception as e:
                    print(f"Fout bij verzenden van e-mails: {str(e)}")
                    # Zelfs als e-mail mislukt, is de offerte opgeslagen.
                    # Overweeg of je hier de transactie wilt terugdraaien of de gebruiker anders wilt informeren.
                    return JsonResponse({
                        "success": False,
                        "error": f"Offerte opgeslagen, maar e-mail versturen mislukt: {str(e)}",
                        "offerte_id": offerte.id # Stuur ID mee zodat er eventueel handmatig actie ondernomen kan worden
                    })

                return JsonResponse({"success": True, "offerte_id": offerte.id})
            else:
                # Formulier is niet valide, stuur de fouten terug
                print(f"Formulier validatie fouten: {form.errors.as_json()}")
                return JsonResponse({
                    "success": False,
                    "errors": json.loads(form.errors.as_json()) # Stuur fouten als JSON
                }, status=400) # Gebruik een 400 status voor client-side errors

        except json.JSONDecodeError:
            return JsonResponse({"success": False, "error": "Ongeldige JSON data"}, status=400)
        except Exception as e:
            print(f"Algemene fout in save_offerte: {str(e)}")
            # Voeg hier meer gedetailleerde logging toe indien nodig (bv. traceback)
            import traceback
            traceback.print_exc()
            return JsonResponse({
                "success": False,
                "error": f"Er is een serverfout opgetreden: {str(e)}"
            }, status=500)

    return JsonResponse({
        "success": False,
        "error": "Ongeldige methode"
    }, status=405) # Method Not Allowed

# klant aanmaken eenvoudigfacturen ----------------------------------------------------------
@login_required
def create_client_eenvoudigfactureren(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            offerte_id = data.get("offerte_id")
            offerte = Offerte.objects.get(id=offerte_id)

            auth_string = f"{settings.EENVOUDIGFACTUREREN_EMAIL}:{settings.EENVOUDIGFACTUREREN_PASSWORD}"
            auth_header = "Basic " + base64.b64encode(auth_string.encode()).decode()
            headers = {
                "Authorization": auth_header,
                "Content-Type": "application/json",
                "Accept": "application/json"
            }

            # Zoek bestaande klanten op basis van e-mail
            search_url = f"https://eenvoudigfactureren.be/api/v1/clients?email_address={offerte.email}"
            search_response = requests.get(search_url, headers=headers)
            
            found_match = False
            if search_response.status_code == 200:
                clients = search_response.json()
                if clients:
                    for client in clients:
                        # Normaliseer adres
                        adres_api = f"{client.get('street', '').strip().lower()}|{client.get('city', '').strip().lower()}"
                        adres_offerte = f"{offerte.straat.strip().lower()} {offerte.nummer.strip().lower()}|{offerte.gemeente.strip().lower()}"

                        if adres_api == adres_offerte:
                            # Exacte match op zowel e-mail als adres gevonden
                            return JsonResponse({
                                "success": True,
                                "client_id": client.get("client_id"),
                                "already_exists": True,
                                "match": "email+adres"
                            })
            
            # Bepaal welke naam te gebruiken
            # Als er een bedrijfsnaam EN BTW-nummer is, gebruik bedrijfsnaam
            # Anders gebruik naam contactpersoon
            if hasattr(offerte, 'naam_bedrijf') and offerte.naam_bedrijf and offerte.btw_nummer:
                client_name = offerte.naam_bedrijf
            else:
                client_name = offerte.naam_contactpersoon
            
            # Geen exacte match op e-mail+adres, dus maak nieuwe klant aan
            url = "https://eenvoudigfactureren.be/api/v1/clients"
            payload = {
                "name": client_name,
                "email_address": offerte.email,
                "street": f"{offerte.straat} {offerte.nummer}",
                "city": offerte.gemeente,
                "country_code": "BE",
                "tax_code": offerte.btw_nummer if offerte.btw_nummer else None,
                "phone_number": offerte.telefoon,
            }

            response = requests.post(url, json=payload, headers=headers)
            if response.status_code == 201:
                return JsonResponse({
                    "success": True,
                    "client_id": response.json().get("client_id"),
                    "already_exists": False
                })
            else:
                return JsonResponse({"success": False, "error": response.text})

        except Offerte.DoesNotExist:
            return JsonResponse({"success": False, "error": "Offerte niet gevonden"})
        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})

    return JsonResponse({"success": False, "error": "Ongeldige methode"})

    


# Alles voor het tonen van de offertes op het dashboard --------------------------------------------------------------
@login_required
def offertes_view(request):
    offertes = Offerte.objects.all().order_by('-created_at')
    return render(request, 'offertes.html', {'offertes': offertes})


@login_required
def delete_offerte(request):
    if request.method == "POST":
        try:
            data = json.loads(request.body)
            offerte_id = data.get("offerte_id")
            offerte = Offerte.objects.get(id=offerte_id)
            offerte.delete()
            return JsonResponse({"success": True})
        except Offerte.DoesNotExist:
            return JsonResponse({"success": False, "error": "Offerte niet gevonden"})
        except Exception as e:
            return JsonResponse({"success": False, "error": str(e)})
    return JsonResponse({"success": False, "error": "Ongeldige methode"})








# Word event overwiew -----------------------------------------------------------------------------------------------
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt, Cm
from django.http import HttpResponse

@login_required
def export_offerte_to_docx(request, offerte_id):
    try:
        # Haal de offerte op uit de database
        offerte = Offerte.objects.get(id=offerte_id)
    except Offerte.DoesNotExist:
        return HttpResponse("Offerte niet gevonden", status=404)

    # Maak een nieuw Word-document
    doc = Document()

    # Stel een kleinere fontgrootte in voor het hele document
    style = doc.styles['Normal']
    font = style.font
    font.size = Pt(10)



    # Maak een tabel voor de klantgegevens met 2 kolommen
    table = doc.add_table(rows=4, cols=4)  # 4 rijen, 4 kolommen (2 paren per rij)
    table.style = 'Table Grid'  # Voeg een raster toe voor zichtbaarheid (optioneel)
    #table.autofit = False  # Schakel autofit uit om handmatige breedtes te gebruiken

    # Stel de kolombreedtes in (in centimeters)
    column_widths = [Cm(4), Cm(6), Cm(4), Cm(6)]  # Label1, Waarde1, Label2, Waarde2
    for col_idx, width in enumerate(column_widths):
        for cell in table.columns[col_idx].cells:
            cell.width = width

    # Definieer de paren voor de tabel
    klantgegevens = [
        ("Naam:", offerte.naam_contactpersoon, "Adres:", offerte.event_adres),
        ("telefoonnr:", offerte.telefoon, "Aantal Personen:", str(offerte.aantal_personen)),
        ("Gewenste Datum:", offerte.gewenste_datum.strftime('%d/%m/%Y'), "Gewenste Tijd:", offerte.gewenste_tijd.strftime('%H:%M')),
        ("Formule:", offerte.formule, "Totaal Prijs:", f"€{offerte.totaal_prijs}"),
    ]

    # Vul de tabel in met extra verticale ruimte
    for row_idx, (label1, waarde1, label2, waarde2) in enumerate(klantgegevens):
        cells = table.rows[row_idx].cells
        # Kolom 1: Label 1 (vetgedrukt)
        p = cells[0].paragraphs[0]
        p.add_run(label1).bold = True
        cells[0].add_paragraph()  # Extra lege paragraaf voor ruimte onder de tekst
        # Kolom 2: Waarde 1
        p = cells[1].paragraphs[0]
        p.add_run(waarde1)
        cells[1].add_paragraph()  # Extra lege paragraaf voor ruimte onder de tekst
        # Kolom 3: Label 2 (vetgedrukt)
        p = cells[2].paragraphs[0]
        p.add_run(label2).bold = True
        cells[2].add_paragraph()  # Extra lege paragraaf voor ruimte onder de tekst
        # Kolom 4: Waarde 2
        p = cells[3].paragraphs[0]
        p.add_run(waarde2)
        cells[3].add_paragraph()  # Extra lege paragraaf voor ruimte onder de tekst

    # Voeg "Extra Info" toe als aparte paragraaf
    p = doc.add_paragraph()
    p.add_run("Extra Info:").bold = True
    p.add_run(f"\t{offerte.extra_info if offerte.extra_info else 'N/A'}")

    # Voeg een lijn toe na de klantgegevens
    doc.add_paragraph("________________________")

    # Voeg de extra vragen toe (zonder schrijflijn)
    extra_vragen = [
        "Zaal /naam: _ _ _ _ _ _ _ _ _ _ _ _",
        "Bakkers: _ _ _ _ _ _ _ _ _ _ _ _",
        "Bakuren: _ _ _ _ _ _ _ _ _ _ _ _",
        "Wagen: _ _ _ _ _ _ _ _ _ _ _ _",
    ]

    for vraag in extra_vragen:
        p = doc.add_paragraph()
        p.add_run(vraag).bold = True

    # Voeg de "Afsluiting" sectie toe met grotere checkboxes
    doc.add_heading('Afsluiting', level=1)
    
    afsluiting_opties = [
        "Wagen volledig leegmaken",
        "Vlees + frieten uitladen, alles omhoog zetten en dweilen",
        "Vlees + frieten en afwas uitladen, alles omhoog zetten en dweilen",
        "Stekker insteken, alles omhoog zetten en dweilen",
    ]

    for optie in afsluiting_opties:
        p = doc.add_paragraph()
        p.add_run("❑ ")  # Groter checkbox symbool
        p.add_run(optie)

    # Voeg de extra tekst toe: "Enkele regels en afspraken"
    doc.add_heading('Enkele regels en afspraken', level=1)
    
    regels = [
        "Filters voor aanvang leegmaken.",
        "Tel het aantal bakjes (zowel friet als snacks) voor en na het bakken.",
        "Zorg voor een verzorgd voorkomen.",
        "Iedereen draagt kledij met het logo van Retrosnacks + schort.",
        "Wees altijd vriendelijk en beleefd.",
        "Respecteer de afgesproken uren.",
        "Informeer de contactpersoon een half uur voor sluiten.",
        "Poets de wagen op rustige momenten.",
    ]

    # Voeg de lijst met opsommingen toe
    for regel in regels:
        doc.add_paragraph(regel, style='List Bullet')

    # Voeg sublijst "Bij thuiskomst" toe
    doc.add_paragraph("Bij thuiskomst:", style='List Bullet')
    sub_regels = [
        "Zorg dat het materiaal proper afgekuist is.",
        "Dweil de wagen.",
    ]
    for sub_regel in sub_regels:
        doc.add_paragraph(sub_regel, style='List Bullet 2')

    # Voeg afsluitende tekst toe
    doc.add_paragraph()
    p = doc.add_paragraph()
    p.add_run("Alvast bedankt voor uw medewerking!").italic = True

    # Maak een HTTP-response met het Word-document
    response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
    response['Content-Disposition'] = f'attachment; filename="offerte_{offerte.id}.docx"'
    doc.save(response)

    return response




# Mail afwijzen event --------------------------------------------------------------
@login_required
def get_offerte_details(request, offerte_id):
    # Ophalen van de offerte via ID
    offerte = get_object_or_404(Offerte, id=offerte_id)
    
    # Data voorbereiden voor JSON-respons
    data = {
        'email': offerte.email,
        'naam_contactpersoon': offerte.naam_contactpersoon,
        'gewenste_datum': offerte.gewenste_datum.strftime('%Y-%m-%d'),
        'gewenste_tijd': offerte.gewenste_tijd.strftime('%H:%M'),
        'naam_bedrijf': offerte.naam_bedrijf,
        'formule': offerte.formule,
        'aantal_personen': offerte.aantal_personen,
        'totaal_prijs': offerte.totaal_prijs,
        'event_adres': offerte.event_adres
    }
    
    # JSON-respons terugsturen
    return JsonResponse(data)





# Status veranderen op dashboard en klant toeveogen -------------------------------------------
@login_required
def update_offerte_status(request, offerte_id):
    if request.method == 'POST':
        offerte = get_object_or_404(Offerte, id=offerte_id)
        nieuwe_status = request.POST.get('status')

        if nieuwe_status:
            offerte.is_bevestigd = nieuwe_status
            offerte.save()

            # === Klant aanmaken als status 'Feest ingepland' is ===
            if nieuwe_status == 'Feest ingepland':
                # Check of klant al bestaat op basis van naam_bedrijf, straat, nummer en gemeente
                klant, aangemaakt = Klant.objects.get_or_create(
                    naam_bedrijf=offerte.naam_bedrijf,
                    straat=offerte.straat,
                    nummer=offerte.nummer,
                    gemeente=offerte.gemeente,
                    defaults={
                        'email': offerte.email,
                        'naam_contactpersoon': offerte.naam_contactpersoon,
                        'telefoon': offerte.telefoon,
                        'btw_nummer': offerte.btw_nummer,
                    }
                )

            return JsonResponse({'success': True, 'message': f'Status gewijzigd naar {nieuwe_status}'})

        return JsonResponse({'success': False, 'message': 'Geen status ontvangen'}, status=400)

    return JsonResponse({'success': False, 'message': 'Alleen POST-verzoeken toegestaan'}, status=405)



# Calender dots events ----------------------------------
@login_required
def kalender_events(request):
    feesten = Offerte.objects.filter(is_bevestigd="Feest ingepland")
    data = [
        {
            "date": feest.gewenste_datum.strftime('%Y-%m-%d'),
            "title": feest.naam_bedrijf,
            "status": feest.is_bevestigd,
            "event_adres": feest.event_adres,
            "gewenste_tijd": feest.gewenste_tijd.strftime('%H:%M') if feest.gewenste_tijd else None,
            "aantal_personen": feest.aantal_personen,
            "formule": feest.formule,
        }
        for feest in feesten
    ]
    return JsonResponse(data, safe=False)



# klanten opslagen ------------------------------------
def export_klanten_excel(request):
    # Create an in-memory output file
    output = io.BytesIO()
    
    # Create Excel workbook and add a worksheet
    workbook = xlsxwriter.Workbook(output)
    worksheet = workbook.add_worksheet('Klanten')
    
    # Add header with bold format
    bold = workbook.add_format({'bold': True})
    headers = ['Naam Contactpersoon', 'Email', 'Straat', 'Nummer', 'Gemeente', 
              'Telefoon', 'Naam Bedrijf', 'BTW Nummer']
    
    for col_num, header in enumerate(headers):
        worksheet.write(0, col_num, header, bold)
    
    # Add data
    klanten = Klant.objects.all()
    for row_num, klant in enumerate(klanten, 1):
        worksheet.write(row_num, 0, klant.naam_contactpersoon)
        worksheet.write(row_num, 1, klant.email)
        worksheet.write(row_num, 2, klant.straat)
        worksheet.write(row_num, 3, klant.nummer)
        worksheet.write(row_num, 4, klant.gemeente)
        worksheet.write(row_num, 5, klant.telefoon)
        worksheet.write(row_num, 6, klant.naam_bedrijf)
        worksheet.write(row_num, 7, klant.btw_nummer or '')
    
    # Auto-fit column width
    for i, header in enumerate(headers):
        worksheet.set_column(i, i, len(header) + 2)
    
    workbook.close()
    
    # Prepare response
    output.seek(0)
    
    # Create HTTP response
    response = HttpResponse(
        output.read(),
        content_type='application/vnd.openxmlformats-officedocument.spreadsheetml.sheet'
    )
    response['Content-Disposition'] = 'attachment; filename=klanten.xlsx'
    
    return response





class ActieListView(ListView):
    model = ActieVanHetMoment
    template_name = 'acties/actie_list.html'
    context_object_name = 'acties'
    ordering = ['-actief_tot']
    paginate_by = 10  # Optioneel: paginering toevoegen

def acties_overzicht(request):
    acties = ActieVanHetMoment.objects.all().order_by('-actief_tot')
    today = date.today()
    
    context = {
        'acties': acties,
        'today': today,
    }
    return render(request, 'acties/acties_overzicht.html', context)